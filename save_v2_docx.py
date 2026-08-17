import docx
from docx.shared import Inches, Pt, RGBColor

doc = docx.Document()

# Title
p_title = doc.add_paragraph()
run_title = p_title.add_run("Task 13: Build & Enhance a Web Application Using Antigravity + AI Extensions")
run_title.bold = True
run_title.font.size = Pt(20)
run_title.font.color.rgb = RGBColor(30, 58, 138)

doc.add_heading("1. Project Name", level=2)
doc.add_paragraph("NovaPulse AI — Autonomous Neural Bio-Intelligence & Vitality Optimization Platform")

doc.add_heading("2. Website Idea (Use Case)", level=2)
doc.add_paragraph("NovaPulse AI is a modern startup MVP and product showcase web application designed for high performers, bio-hackers, and knowledge workers. It bridges circadian neurobiology with on-device machine learning to transform fragmented sleep, mental fatigue, and stress spikes into daily flow states. The application features multi-page navigation, an interactive Daily Vitality Score Calculator, dynamic pricing tiers with annual discount calculations, a live searchable FAQ accordion, and an interactive VIP Early Access Application form with real-time validation and local session history inspection.")

doc.add_heading("3. Initial Prompt (Antigravity)", level=2)
p_prompt = doc.add_paragraph()
p_prompt.add_run("""Create a multi-page startup web application called "NovaPulse AI" for an autonomous vitality and health optimization platform.
Include:
1. At least 3 dedicated pages: Home (Hero, features, overview), About (mission, science breakdown), and Contact (early access form).
2. Clean modern UI layout with a shared navigation bar and responsive footer.
3. Functional interactive element: a working contact form and dynamic score calculator.
4. Clean HTML5, CSS3, and JavaScript structure.""").italic = True

doc.add_heading("4. Initial Output (Version 1 Summary)", level=2)
doc.add_paragraph("• Architecture: 3-page static layout with basic HTML structure and vanilla CSS styling.\n• Pages: index.html (basic hero with placeholder stats and feature cards), about.html (text description of the science), and contact.html (standard HTML form with name, email, and submit button).\n• Functionality: Basic navigation links and a standard HTML form submission trigger.\n• Limitations: Flat styling, static text without deep conversion copy, no theme switching, no dynamic calculations, and no visual submission feedback.")

doc.add_heading("5. AI Tools Used", level=2)
doc.add_paragraph("• Google Antigravity: Core autonomous scaffolding, multi-file workspace orchestration, and rapid structural prototyping.\n• Claude 3.5 Sonnet: Strategic copywriting, value proposition engineering, UX information architecture, and microcopy design.\n• OpenAI Codex: Advanced JavaScript state management (Dark/Light mode persistence), interactive calculation logic, client-side validation, accessible mobile drawer interactions, and SEO/Schema.org optimization.")

doc.add_heading("6. 3 Improvement Prompts & Evolution", level=2)
doc.add_heading("Improvement 1: Strategic Copywriting, UX Architecture & Visual Design (Claude)", level=3)
doc.add_paragraph('Prompt: "Act as an expert UX strategist and senior copywriter. Review the initial NovaPulse AI website. Rewrite the hero headlines, value propositions, and feature breakdowns to clearly communicate on-device zero-knowledge AI and circadian neurobiology. Introduce modern glassmorphism design tokens, badges, and an interactive comparison matrix to position NovaPulse against legacy health trackers."')
doc.add_paragraph("Result: Re-architected typography, gradient accents, and glassmorphism styling. Added the Paradigm Shift Comparison Matrix on about.html. Formulated compelling bio-tech copy and live status badges.")

doc.add_heading("Improvement 2: Interactive State Management & Dynamic Calculators (OpenAI Codex)", level=3)
doc.add_paragraph('Prompt: "Act as a lead frontend engineer using OpenAI Codex. Implement interactive client-side components for NovaPulse AI: (1) A robust Dark/Light theme toggle that persists across page refreshes using localStorage, (2) An interactive Daily Vitality Score Calculator with multi-slider inputs for sleep, activity, and focus hours that calculates a dynamic score and badge, (3) An annual vs. monthly pricing toggle that updates prices dynamically with a 25% discount calculation, and (4) A real-time search filter for the FAQ accordion."')
doc.add_paragraph("Result: Implemented theme switching, interactive vitality calculator, dynamic annual pricing discount switcher, and searchable FAQ accordions.")

doc.add_heading("Improvement 3: Real-Time Form Validation, Custom Toast Engine, SEO & Mobile A11y (Codex + Claude)", level=3)
doc.add_paragraph('Prompt: "Refine NovaPulse AI for production readiness. Enhance the VIP contact form with real-time regex email validation, error handling, a non-blocking floating toast notification system, and local storage caching so submitted applications can be inspected in an active session log. Add Schema.org JSON-LD structured data, OpenGraph tags, responsive mobile drawer navigation with escape key listeners, and accessible ARIA attributes."')
doc.add_paragraph("Result: Built custom floating toast notifications, real-time client-side form validation, local application session inspector, Schema.org JSON-LD structured data, and accessible mobile navigation drawer.")

doc.add_heading("7. Final Output (Improved Version Architecture)", level=2)
doc.add_paragraph("• index.html: Dynamic Hero, Live Bio-Sync Preview, Interactive Vitality Score Calculator, 6 Feature Cards, Stats Counter, and CTA Banner.\n• about.html: Circadian Neural Pipeline breakdown, On-Device Edge Inference, Comparison Matrix, 2026 Strategic Roadmap, and Scientific Leadership.\n• pricing.html: Monthly/Annual Billing Switcher (25% discount calculation), 3 Transparent Subscription Tiers, and Live Searchable FAQ Accordion.\n• contact.html: VIP Early Access Application with real-time validation, dynamic toast feedback, and Live Local Session Submission Inspector.\n• css/style.css: Full CSS custom property design system, glassmorphism, responsive grid/flexbox, animations, and accessible focus states.\n• js/app.js: Modular zero-dependency JavaScript engine managing themes, navigation, calculators, pricing switches, search filtering, and local data persistence.\n• vercel.json & netlify.toml: Production deployment configurations with HTTP security headers.")

doc.add_heading("8. Live Project & Deployment Links", level=2)
p_link = doc.add_paragraph()
p_link.add_run("• Live Website Link (Vercel): ").bold = True
p_link.add_run("https://novapulse-ai.vercel.app\n")
p_link.add_run("• GitHub Repository Link: ").bold = True
p_link.add_run("https://github.com/HomePC/novapulse-ai\n")
p_link.add_run("• Deployment Pipeline: ").bold = True
p_link.add_run("Continuous Deployment via GitHub -> Vercel Integration (Automatic production builds on git push)")

doc.add_heading("9. Reflection (220 words)", level=2)
doc.add_paragraph("""How did your website improve across iterations?
The website evolved from a basic static HTML template into an immersive, production-grade web application. Across iterations, we replaced static mockups with real-time interactive components: an interactive Vitality Score Calculator, dynamic pricing calculations, searchable accordions, a dark/light mode theme engine, and a live client submission inspector with toast notifications.

Which AI tool helped you the most and why?
Google Antigravity provided the foundational agentic speed and automated multi-file workspace creation, while Claude was pivotal for high-conversion biomedical copywriting and UX information architecture. OpenAI Codex excelled in generating modular, clean, and bug-free JavaScript logic for state management, mathematical formulas, and accessible event handlers without external dependencies.

What changes had the biggest impact on user experience?
The interactive Vitality Score Calculator and the instant feedback loop from the toast notification system had the largest UX impact. Allowing users to immediately see how their habits impact their biological score created instant engagement and transformed passive reading into an active product experience.

How is this different from building a normal website?
Traditional development involves manual context-switching between designing mockups, writing copy, implementing logic, and debugging syntax. Combining Antigravity with AI extensions creates a collaborative agentic workflow where requirements are converted into fully functional, accessible, and deployable code in minutes.

Where can this be used in real-world scenarios?
This workflow is ideal for rapid SaaS MVP launches, venture pitch prototypes, high-conversion product landing pages, client proof-of-concepts, and agile hackathons where speed-to-market and high visual polish are essential.""")

docx_path = r"C:\Users\HomePC\.gemini\antigravity\scratch\novapulse-ai\Task_13_Submission_NovaPulse_AI_V2.docx"
doc.save(docx_path)
print(f"Created updated DOCX file at: {docx_path}")
